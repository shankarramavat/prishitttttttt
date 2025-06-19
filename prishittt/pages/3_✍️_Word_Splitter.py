import streamlit as st
import docx
import zipfile
import tempfile
import os
import io

# --- Backend function is correct, no changes needed ---
def split_and_zip_word(uploaded_file, delimiter, log_container):
    try:
        source_doc = docx.Document(uploaded_file)
        log_container.info("✅ Word document loaded successfully.")
    except Exception as e:
        log_container.error(f"❌ Could not read Word file. Details: {e}")
        return None, 0
        
    files_created = 0
    doc_count = 0
    current_doc = None
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        with tempfile.TemporaryDirectory() as temp_dir:
            for para in source_doc.paragraphs:
                if delimiter in para.text:
                    if current_doc is not None:
                        filename = f"Document_{doc_count}.docx"
                        output_path = os.path.join(temp_dir, filename)
                        current_doc.save(output_path)
                        zipf.write(output_path, arcname=filename)
                        files_created += 1
                        log_container.write(f"  - Saved and added `{filename}`")
                    doc_count += 1
                    current_doc = docx.Document()
                if current_doc is not None:
                    current_doc.add_paragraph(para.text)
            if current_doc is not None:
                filename = f"Document_{doc_count}.docx"
                output_path = os.path.join(temp_dir, filename)
                current_doc.save(output_path)
                zipf.write(output_path, arcname=filename)
                files_created += 1
                log_container.write(f"  - Saved and added `{filename}`")

    if files_created == 0:
        log_container.warning("⚠️ No documents were created. The delimiter text was not found.")
        return None, 0
    log_container.success(f"✅ Found and created {files_created} documents.")
    return zip_buffer.getvalue(), files_created

# --- UI LOGIC (CORRECTED) ---
st.set_page_config(page_title="Word Splitter", layout="wide")
st.title("✍️ Word (DOCX) Splitter")

with st.container(border=True):
    st.header("1. Configure Split Settings")
    delimiter = st.text_input("Enter the text that marks the start of a new document (e.g., 'Client Statement')")
    st.info("Note: This basic splitter works on paragraph text. Tables, images, and complex formatting may not be carried over.", icon="⚠️")

with st.container(border=True):
    st.header("2. Upload Your Word File")
    uploaded_file = st.file_uploader("Upload a consolidated .docx file", type="docx", label_visibility="collapsed")

if st.button("🚀 Process Word File", type="primary", use_container_width=True):
    if uploaded_file is not None and delimiter:
        st.header("3. Results")
        with st.spinner("Splitting Word document..."):
            log_expander = st.expander("Show Processing Logs", expanded=True)
            zip_data, files_created = split_and_zip_word(uploaded_file, delimiter, log_expander)
            if zip_data:
                st.metric("Files Created", files_created)
                st.download_button(
                    label="⬇️ Download ZIP File",
                    data=zip_data,
                    file_name=f"Split_{uploaded_file.name.replace('.docx', '.zip')}",
                    mime="application/zip",
                    use_container_width=True
                )
    else:
        st.warning("⚠️ Please upload a file and enter the delimiter text.")